"""
Processor for text-based files: TXT, DOCX, MARKDOWN, etc.
Converts them to normalized Markdown format.
"""

import io
import logging

from mkb.processors.base import ProcessingResult, TextualProcessor

logger = logging.getLogger(__name__)


class TextProcessor(TextualProcessor):
    """Handles plain text and simple document formats."""
    
    supported_mime_types = [
        "text/plain",
        "text/markdown",
        "text/x-markdown",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]
    
    def can_process(self, mime_type: str, filename: str) -> bool:
        lower_name = filename.lower()
        return (
            mime_type in self.supported_mime_types
            or lower_name.endswith((".txt", ".md", ".markdown", ".docx"))
            or mime_type.startswith("text/")
        )
    
    def process(self, data: bytes, filename: str) -> ProcessingResult:
        """Convert text or document to Markdown."""
        try:
            lower_name = filename.lower()
            
            if lower_name.endswith(".docx"):
                return self._process_docx(data)
            else:
                # Plain text is already markdown-compatible
                return self._process_plain_text(data, filename)
        except Exception as e:
            return ProcessingResult(
                processing_type=self.get_processing_type(),
                output_format="md",
                content=b"",
                error=f"Text processing failed: {str(e)}"
            )
    
    def _process_plain_text(self, data: bytes, filename: str) -> ProcessingResult:
        """Process plain text files."""
        try:
            # Try UTF-8 first
            text = data.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            text = data.decode('latin-1', errors='replace')
        
        # Add filename as header if not markdown
        if not filename.lower().endswith(".md"):
            text = f"# {filename}\n\n{text}"
        
        return ProcessingResult(
            processing_type=self.get_processing_type(),
            output_format="md",
            content=text.encode('utf-8'),
            conversion_metadata={
                "method": "plain_text",
                "detected_encoding": "utf-8",
            }
        )
    
    def _process_docx(self, data: bytes) -> ProcessingResult:
        """Process DOCX files using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )
        
        # Parse DOCX from bytes
        docx_file = io.BytesIO(data)
        doc = Document(docx_file)
        
        # Extract text while preserving structure
        markdown_parts = []
        
        for para in doc.paragraphs:
            # Preserve heading style
            if para.style.name.startswith('Heading'):
                level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                markdown_parts.append(f"{'#' * level} {para.text}")
            else:
                markdown_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            markdown_parts.append("\n")
            markdown_parts.append(self._table_to_markdown(table))
        
        markdown_content = "\n\n".join(filter(None, markdown_parts))
        
        return ProcessingResult(
            processing_type=self.get_processing_type(),
            output_format="md",
            content=markdown_content.encode('utf-8'),
            conversion_metadata={
                "method": "docx",
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
            }
        )
    
    @staticmethod
    def _table_to_markdown(table) -> str:
        """Convert DOCX table to Markdown."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        
        if len(rows) > 1:
            # Add separator after header
            rows.insert(1, "|" + "|".join("---" for _ in rows[0].split("|")[1:-1]) + "|")
        
        return "\n".join(rows)
